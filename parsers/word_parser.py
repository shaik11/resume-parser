from docx import Document
from parsers.base import FileParser

class WordParser(FileParser):
    """
    Extracts text from WORD file resumes using python-docx.
    """

    def parse(self, file_path: str) -> str:
        extracted_text = []

        try:
            document = Document(file_path)

            # Extract paragraph text
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text.strip())

            # Extract table text (important for resumes)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            extracted_text.append(cell_text)

        except Exception as exc:
            raise RuntimeError(f"Failed to parse Word document: {exc}") from exc
        # print("\n".join(extracted_text))
        return "\n".join(extracted_text)